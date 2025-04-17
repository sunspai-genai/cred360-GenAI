import json
import logging
import os
import queue
import re

# from sqlparse.sql import Token, TokenList, Where, Comparison, Identifier
from collections import defaultdict
from datetime import datetime
from logging.handlers import TimedRotatingFileHandler
from pathlib import Path
from typing import Any, Dict, Iterator, List, Optional, Set, Union

# Logging
import coloredlogs
import docx
import markdown
import ollama
import openai
import pandas as pd
import sqlparse
import streamlit as st
import torch
from bs4 import BeautifulSoup
from docx.shared import Inches, Pt, RGBColor
from langchain.docstore.document import Document
from langchain.document_loaders.base import BaseLoader
from langchain_community.embeddings.fastembed import FastEmbedEmbeddings
from langchain_huggingface import HuggingFaceEmbeddings

# from langchain_community.chat_models import ChatOllama
from langchain_ollama import ChatOllama, OllamaEmbeddings
from langchain_openai import AzureChatOpenAI, ChatOpenAI
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils.dataframe import dataframe_to_rows
from sqlglot import Expression, exp, parse_one
from sqlglot.optimizer import optimize
from streamlit.logger import get_logger
from streamlit.runtime.scriptrunner import add_script_run_ctx

logger = logging.getLogger('cred360')
logger.handlers = []  # Clear existing handlers
log_format_string = '%(asctime)s %(pathname)s: %(levelname)s: %(funcName)s(): %(lineno)d: %(message)s'
formatter = logging.Formatter(log_format_string)
level_styles = {
    'info': {'color': 'magenta'},
    'debug': {'color': 'green'},
    'warning': {'color': 'yellow'},
    'error': {'color': 'red'},
    'critical': {'color': 'red', 'bold': True},}
coloredlogs.install(level='DEBUG', logger=logger, fmt=log_format_string, level_styles=level_styles)

# file_handler = TimedRotatingFileHandler(".\\logs\\, when="midnight", interval=1, backupCount=100)
file_handler = TimedRotatingFileHandler(filename="logs/" + "jedi.log", when="midnight", interval=1, backupCount=100)
file_handler.setFormatter(formatter)
logger.addHandler(file_handler)

console_handler = logging.StreamHandler()
console_handler.setFormatter(formatter)
logger.addHandler(console_handler)

logging.getLogger('LiteLLM').setLevel(logging.ERROR)
logging.getLogger('httpcore').setLevel(logging.ERROR)
logging.getLogger('httpx').setLevel(logging.ERROR)
logging.getLogger('backoff').setLevel(logging.ERROR)
logging.getLogger('opentelemetry').setLevel(logging.ERROR)

class BaseDataFrameLoader(BaseLoader):
    def __init__(self, data_frame: Any, *, page_content_column: Union[str, List[str]] = "text"):
        """Initialize with dataframe object.
        Args:
            data_frame: DataFrame object.
            page_content_column: Name of the column or list of column names containing the page content.
            Defaults to "text".
        """
        self.data_frame = data_frame
        self.page_content_column = page_content_column

    def lazy_load(self) -> Iterator[Document]:
        """Lazy load records from dataframe.""" 
        for _, row in self.data_frame.iterrows():
            if isinstance(self.page_content_column, list):
                text = ' '.join(f'{col}:{row[col]}' for col in self.page_content_column)
            else:
                # If it's a single string, use the column name directly
                # text = f'{col}:{row[self.page_content_column]}'
                text = f'{self.page_content_column}:{row[self.page_content_column]}'

            metadata = row.to_dict()
            if isinstance(self.page_content_column, list):
                for col in self.page_content_column:
                    metadata.pop(col, None)
            else:
                metadata.pop(self.page_content_column, None)
            yield Document(page_content=text, metadata=metadata)
 
    def load(self) -> List[Document]:
        """Load full dataframe."""
        return list(self.lazy_load())

class DataFrameLoader(BaseDataFrameLoader):
    """Load `Pandas` DataFrame.""" 
    def __init__(self, data_frame: Any, page_content_column: Union[str, List[str]] = "text"):
        """Initialize with dataframe object.
        Args:
            data_frame: Pandas DataFrame object.
            page_content_column: Name of the column or list of column names containing the page content.
            Defaults to "text".
        """
        try:
            import pandas as pd
        except ImportError as e:
            raise ImportError(
                "Unable to import pandas, please install with `pip install pandas`."
            ) from e

        if not isinstance(data_frame, pd.DataFrame):
            raise ValueError(
                f"Expected data_frame to be a pd.DataFrame, got {type(data_frame)}"
            )
        super().__init__(data_frame, page_content_column=page_content_column)





@st.cache_resource
def configure_embedding_model():
    # Detect if a CUDA-enabled GPU is available
    if torch.cuda.is_available():
        device = 'cuda'
    else:
        device = 'cpu'
    model_name = os.getenv("EMBEDDING_MODEL")
    logger.debug(f"The Embedding Model Name is: {model_name}")
    # model_kwargs = {'device': device}
    # encode_kwargs = {'normalize_embeddings': False}
    embedding_model = OllamaEmbeddings(model=model_name)
    return embedding_model



def markdown_to_word(markdown_text, output_file):
    # Convert Markdown to HTML
    html = markdown.markdown(markdown_text)
    # Parse the HTML
    soup = BeautifulSoup(html, 'html.parser')
    # Create a new Word document
    doc = docx.Document()
    # Add content to the Word document
    for element in soup.descendants:
        print(element)
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='ListBullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='ListNumber')
        elif element.name == 'strong':
            doc.add_paragraph(element.get_text(), style='Normal')
        elif element.name == 'em':
            doc.add_paragraph(element.get_text(), style='Emphasis')
        elif element.name == 'img':
            # Get the image URL
            img_url = element['src']
            # print(img_url)
            doc.add_picture(img_url, width=Inches(5.0)) 
        elif element.name == 'table':
            rows = element.find_all('tr')
            word_table = doc.add_table(rows=0, cols=len(rows[0].find_all(['th', 'td'])))
            for row in rows:
                cells = row.find_all(['th', 'td'])
                row_cells = word_table.add_row().cells
                for idx, cell in enumerate(cells):
                    row_cells[idx].text = cell.get_text(strip=True)

    # Save the Word document
    doc.save(output_file)
    # Save the document to a buffer
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


###################### Document Processing Class #######################################